from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 1. Monthly Progress Report
doc = Document()
doc.styles['Normal'].font.name = 'Arial'
doc.styles['Normal'].font.size = Pt(11)

title = doc.add_paragraph()
title_run = title.add_run('Monthly Progress Report')
title_run.bold = True
title_run.font.size = Pt(12)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('(This has to be approved and signed by the Placement Manager at the end of every month )\n', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER

table1 = doc.add_table(rows=2, cols=3)
table1.style = 'Table Grid'
# Row 0
table1.cell(0, 0).text = 'Name:\n\nReg. No :'
table1.cell(0, 1).text = 'Date'
table1.cell(0, 2).text = 'Progress Report No...........................'
# Row 1
table1.cell(1, 0).text = 'Week'
table1.cell(1, 1).text = 'Brief Description of Work Carried Out'
table1.cell(1, 2).text = '' # Empty col

# Add 4 empty rows for the weeks
for i in range(4):
    row = table1.add_row()

doc.add_paragraph()

table2 = doc.add_table(rows=2, cols=1)
table2.style = 'Table Grid'
table2.cell(0, 0).text = 'Problems Encountered and Solutions Found'
table2.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
# Make second row larger
row2 = table2.rows[1]
row2.height = Inches(2)

doc.add_paragraph()
doc.add_paragraph('Leave Record:').bold = True
doc.add_paragraph('No of days absent in this month: _________________')

table3 = doc.add_table(rows=1, cols=2)
table3.style = 'Table Grid'
table3.cell(0, 0).text = 'Approved'
table3.cell(0, 1).text = 'Not Approved'

doc.add_paragraph('\nPlacement Manager\'s Name :')
doc.add_paragraph('\n..............................                                              ..............................')
doc.add_paragraph('Signature                                                   Date')

doc.save('public/templates/Monthly_Progress_Report.docx')

# 2. Weekly Internship Log
doc2 = Document()
doc2.styles['Normal'].font.name = 'Arial'
doc2.styles['Normal'].font.size = Pt(11)

table1 = doc2.add_table(rows=1, cols=2)
table1.style = 'Table Grid'
table1.cell(0, 0).text = 'Name:'
table1.cell(0, 1).text = 'Week Ending..........................'
table1.rows[0].height = Inches(0.5)

doc2.add_paragraph()

table2 = doc2.add_table(rows=8, cols=3)
table2.style = 'Table Grid'
table2.cell(0, 0).text = 'Day'
table2.cell(0, 1).text = 'Brief Description of Work Carried Out'
table2.cell(0, 2).text = ''

days = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']
for idx, day in enumerate(days):
    table2.cell(idx + 1, 0).text = day
    table2.rows[idx + 1].height = Inches(0.5)

doc2.add_paragraph()

table3 = doc2.add_table(rows=2, cols=1)
table3.style = 'Table Grid'
table3.cell(0, 0).text = 'Problems Encountered and Solutions Found'
table3.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
table3.rows[1].height = Inches(2)

doc2.save('public/templates/Weekly_Log_Template.docx')

print("Templates generated.")
